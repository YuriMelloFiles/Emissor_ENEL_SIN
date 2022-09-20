#!/usr/bin/env python
# coding: utf-8

# In[61]:


from docx2pdf import convert
import pyproj
import pickle
import locale
from openpyxl import drawing
from pycep_correios import get_address_from_cep, WebService, exceptions
import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from functools import partial
from PIL import ImageTk, Image
from tkinter import filedialog
import tkinter.font as font
from tkinter.messagebox import showinfo
from docx.shared import RGBColor
import os
from os import path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
from os.path import expanduser
from datetime import date, timedelta
import datetime
from openpyxl import load_workbook 
root = Tk()
root.title('Emissor de documentos Enel')

#---Estrutura para encontrar caminho do ARQUIVO/ICONE
home = expanduser("~")
for r,d,f in os.walk(home + "\DEV SINERGIA"):
        for files in f:
                if files == "Python_ENEL_MD_SE_Base_16092022.docx":
                    caminho = caminhoIcone = os.path.join(r,files)
caminhoBackup = caminho.replace("\Python_ENEL_MD_SE_Base_16092022.docx", "\\backup.txt")                     
icone = caminhoIcone.replace("\Python_ENEL_MD_SE_Base_16092022.docx", "\ICONE.ico") 
#---Iniciando tela tkinter
root.iconbitmap(icone) 
myFont = font.Font(family='Helvetica', size=10, weight='bold')
root.geometry('1020x660')
root['background'] = '#ff4787'
root.resizable(width=False, height=False)
messagebox.showinfo(message="Seja bem vindo(a) ao nosso gerador de documentos Enel !Aqui vão algumas dicas: \n\n ● Nos campos de endereço, preencha o CEP primeiro. Assim você terá uma ajudinha extra :) \n\n ● No campo de tensão secundária do transformador, deverá ser preenchino no formato tensão de linha / tensão de fase (zzz/xxx)\n\n ● Para representar casas decimais, use ponto. \n\n ● Ao preenhcer o campo de 'Descrição da atividade' , certifique-se que não há linhas em branco. Elas podem comprometer a formatação do documento.")

#---Criando scrollbar

#Criando main Frame
main_frame = Frame(root)
main_frame.pack(fill=BOTH , expand=1)
#Criando a Canva
my_canvas = Canvas(main_frame)
my_canvas.pack(side=LEFT , fill=BOTH , expand=1)
#Adicionando a Scrollbar no Canvas
my_scrollbar = ttk.Scrollbar(main_frame , orient=VERTICAL , command=my_canvas.yview)
my_scrollbar.pack(side=RIGHT,fill=Y)
#Configurando Canvas
my_canvas.configure(yscrollcommand=my_scrollbar.set)
my_canvas.bind('<Configure>', lambda e: my_canvas.configure(scrollregion = my_canvas.bbox("all")))
my_canvas.bind_all('<MouseWheel>', lambda event: my_canvas.yview_scroll(int(-1*(event.delta/120)), "units"))
#Criando outro frame dentro do Canvas
second_frame = Frame(my_canvas)
second_frame['background'] = 'white'
my_canvas['background'] = 'white'
my_canvas.create_window((0,0), window=second_frame , anchor="nw" )



#Funcoes  
_projections = {}
def zone(coordinates):
        if 56 <= coordinates[0] < 64 and 3 <= coordinates[1] < 12:
            return 32
        if 72 <= coordinates[0] < 84 and 0 <= coordinates[1] < 42:
            if coordinates[1] < 9:
                return 31
            elif coordinates[1] < 21:
                return 33
            elif coordinates[1] < 33:
                return 35
            return 37
        return int((coordinates[1] + 180) / 6) + 1

def letter(coordinates):
        return 'CDEFGHJKLMNPQRSTUVWXX'[int((coordinates[0] + 80) / 8)]

def project(coordinates):
        z = zone(coordinates)
        l = letter(coordinates)
        if z not in _projections:
            _projections[z] = pyproj.Proj(proj='utm', zone=z, ellps='WGS84')
        x, y = _projections[z](coordinates[1], coordinates[0])
        if y < 0:
            y += 10000000
        return z, l, x, y

def unproject(z, l, x, y):
        if z not in _projections:
            _projections[z] = pyproj.Proj(proj='utm', zone=z, ellps='WGS84')
        if l < 'N':
            y -= 10000000
        lng, lat = _projections[z](x, y, inverse=True)
        return (lng, lat)
#Função para checar se os campos estão preenchidos
def checkVariaveis(a,b,c1,c2,d,e,f,g,h,i,j):

    if len(a.get()) == 0 or len(b.get()) == 0 :
        messagebox.showinfo(message="É necessário preencher as dimensões do módulo!")
        return 0
    if len(c1.get()) == 0:
        messagebox.showinfo(message="É necessário preencher o CEP da usina!")
        return 0
    if len(c2.get()) == 0:
        messagebox.showinfo(message="É necessário preencher o CEP do titular!")
        return 0
    if len(d.get()) == 0 or len(e.get()) == 0:
        messagebox.showinfo(message="É necessário preencher Quantidade de inversores, potência do inversor!")
        return 0
    if len(f.get()) == 0 or len(g.get()) == 0:
        messagebox.showinfo(message="É necessário preencher Quantidade de módulos, potência do módulo!")
        return 0
    if len(h.get()) == 0 or len(i.get()) == 0:
        messagebox.showinfo(message="É necessário preencher Quantidade de transformadores, potência do transformador!")
        return 0
    if len(j.get()) == 0:
        messagebox.showinfo(message="É necessário preencher a tensão secundária do transformador!")
        return 0
    
def deg_to_dms(deg, pretty_print=None, ndp=4):
    m, s = divmod(abs(deg)*3600, 60)
    d, m = divmod(m, 60)
    if deg < 0:
        d = -d
    d, m = int(d), int(m)

    if pretty_print:
        if pretty_print=='latitude':
            hemi = 'N' if d>=0 else 'S'
        elif pretty_print=='longitude':
            hemi = 'L' if d>=0 else 'O'
        else:
            hemi = '?'
        return '{d:d}° {m:d}′ {s:.{ndp:d}f}″ {hemi:1s}'.format(
                    d=abs(d), m=m, s=s, hemi=hemi, ndp=ndp)
    return d, m, s

def backup():
    
    with open(caminhoBackup, 'rb') as f:
        dict = pickle.load(f)
    NomeUsina.insert(0,dict[2])
    TitularUc.insert(0,dict[4])
    CNPJ.insert(0,dict[5])
    ART.insert(0,dict[27])
    LocalizacaoEnel.insert(0,dict[11])
    
    ModeloModulo.insert(0,dict[12])
    PotenciaModulo.insert(0,dict[13])
    QuantidadeTotalModulos.insert(0,dict[14])
    ModeloInversor.insert(0,dict[15])
    PotenciaInversor.insert(0,dict[16])
    QuantidadeInversor.insert(0,dict[17])
    TipoDeEstrutura.insert(0,dict[18])
    FabricanteModulo.insert(0,dict[19])
    FabricanteInversor.insert(0,dict[20])
    PotenciaTransformador.insert(0,dict[21])
    TipoTransformador.insert(0,dict[22])
    TensaoSecundario.insert(0,dict[23])
    AlturaModulo.insert(0,dict[25])
    LarguraModulo.insert(0,dict[26])
    
    
    TelefoneTitular.insert(0,dict[33])
    EmailTitular.insert(0,dict[34])
    DescricaoAtividade.insert(1.0,dict[35])
    CodigoCNAE.insert(0,dict[36])
    Rele.insert(0,dict[37])
    FabricanteRele.insert(0,dict[38])
    TC.insert(0,dict[39])
    FabricanteTC.insert(0,dict[40])
    NivelIsolacao.insert(0,dict[41])
    CorrenteSaidaInversor.insert(0,dict[42])
    
    QuantidadeTransformadores.insert(0,dict[43])
    Impedancia.insert(0,dict[44])
    InscricaoEstadual.insert(0,dict[45])
    InscricaoMunicipal.insert(0,dict[46])
    
    RepresentanteLegalUm.insert(0,dict[47])
    CPFRepresentanteLegalUm.insert(0,dict[48])
    RGRepresentanteLegalUm.insert(0,dict[49])
    OrgaoEmissorRGUm.insert(0,dict[50])
    CargoUm.insert(0,dict[51])
    CotaUm.insert(0,dict[52])
    TelefoneCelularRepresentanteLegalUm.insert(0,dict[53])
    EmailRepresentanteLegalUm.insert(0,dict[54])
    RepresentanteLegalDois.insert(0,dict[55])
    CPFRepresentanteLegalDois.insert(0,dict[56])
    RGRepresentanteLegalDois.insert(0,dict[57])
    OrgaoEmissorRGDois.insert(0,dict[58])
    CargoDois.insert(0,dict[59])
    CotaDois.insert(0,dict[60])
    TelefoneCelularRepresentanteLegalDois.insert(0,dict[61])
    EmailRepresentanteLegalDois.insert(0,dict[62])
    PrevisaoEnergizacao.insert(0,dict[63])

#Função para criar os documentos    

def imp():
    
    messagebox.showinfo(message="Executando o programa.. \n\n Isso pode levar alguns segundos ⏳")
    TextBox=DescricaoAtividade.get("1.0","end")
    checkVariaveis(AlturaModulo,LarguraModulo,CepEnel,CepTitular,QuantidadeInversor,PotenciaInversor,QuantidadeTotalModulos,PotenciaModulo,PotenciaTransformador,QuantidadeTransformadores,TensaoSecundario)

    QuantidadeKwp = str((int(QuantidadeTotalModulos.get()) * int(PotenciaModulo.get()))/1000) 
    QuantidadeKwca = str((int(QuantidadeInversor.get()) * int(PotenciaInversor.get()))) 
    QuantidadeModulosPorInversor = str((int(QuantidadeTotalModulos.get()) // int(QuantidadeInversor.get()))) 
    PotenciaCCPorInversor = str(((int(QuantidadeTotalModulos.get()) // int(QuantidadeInversor.get())) * int(PotenciaModulo.get())/1000) )
    AreaTotal = str(round((float(AlturaModulo.get()) * float(LarguraModulo.get()) *float(QuantidadeTotalModulos.get())),2))
    PotenciaTotalTransformadores = str((int(QuantidadeTransformadores.get()) * int(PotenciaTransformador.get())))
    QCPotAtiva = str(float(QuantidadeKwca) + 3.4)
    QCPotAparente = str(float(QuantidadeKwca) + 4.25)
    
    #convertendo coordenada para usar no RCG
    Loc = LocalizacaoEnel.get()
    chars = '°º'
    Loc = Loc.translate(str.maketrans('', '', chars))
    Loc = Loc.split(",")
    Loc = str(deg_to_dms(float(Loc[0]), pretty_print='latitude')) +" , " + '\n'+  str(deg_to_dms(float(Loc[1]), pretty_print='longitude'))
    CoordenadasGMSParaRCG = Loc
    
    #convertendo coordenadas para UTM
    LocUTM = LocalizacaoEnel.get()
    chars = '°º'
    LocUTM = LocUTM.translate(str.maketrans('', '', chars))
    LocUTM = LocUTM.split(",")
    coordinate = [float(LocUTM[0]) , float(LocUTM[1])]
    z,l,x,y = project(coordinate)
    lng,lat = unproject(z,l,x,y)
    crd=[lat,lng]
    a,b,c,d = project(crd)
    CoordenadasUTM = str(round(c,2)) + " m E ,  \n" + str(round(d,2)) + " m S" + " ,  Zona: " + str(a) + str(b)
    data_atual = date.today()
    dataAtual = data_atual.strftime('%d/%m/%Y')
    
    
    document = Document(caminho) #deve ser variavel caminho
     #creating style
    styles = document.styles
    style = styles.add_style('titulo', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial'
    font.size = Pt(18)
    font.bold = True
    font.color.rgb = RGBColor(19, 40,105)
    
    
    
    CorrenteCalculadaTransformador = (float(13.8)/(12 * 1.732 ))
    CorrenteCalculadaTransformador = str(round(CorrenteCalculadaTransformador,1))
    
    #Tratamento de endereço dinâmico
    
    if len(CidadeEnel.get())==0:
        CidadeEnel.set(address['cidade'])
    
    if len(RuaEnel.get())==0:
        RuaEnel.set(address['logradouro'])
        
    if len(BairroEnel.get())==0:
        BairroEnel.set(address['bairro'])
        
        
    if len(CidadeTitular.get())==0:
        CidadeTitular.set(addressTitular['cidade'])
    
    if len(RuaTitular.get())==0:
        RuaTitular.set(addressTitular['logradouro'])
        
    if len(BairroTitular.get())==0:
        BairroTitular.set(addressTitular['bairro'])
    
    if len(UfTitular.get())==0:
        UfTitular.set(addressTitular['uf'])
    
    
    if (TipoTransformador.get()).upper() == "A SECO":
        CorrenteInrush = "14"
    else:
        CorrenteInrush = "8"
    
    
    
    if Uc.get() == 1:
        NumeroUc_final = "Uc não existente"
    if Uc.get() != 1 :
        NumeroUc_final = NumeroUc.get()
        
        


    for p in document.paragraphs:
        
        if '$PotenciaTotalTransformadores' in p.text:
            p.text = p.text.replace('$PotenciaTotalTransformadores', PotenciaTotalTransformadores)
            
        if '$Rele' in p.text:
            p.text = p.text.replace('$Rele', rele.get())
        if '$FabricanteRele' in p.text:
            p.text = p.text.replace('$FabricanteRele', FabricanteRele.get())
            
        if '$QuantidadeTransformadores' in p.text:
            p.text = p.text.replace('$QuantidadeTransformadores', QuantidadeTransformadores.get())
        if '$CidadeEnel' in p.text:
            p.text = p.text.replace('$CidadeEnel', CidadeEnel.get())
        if '$NumProjeto' in p.text:
            p.text = p.text.replace('$NumProjeto', NumProjeto.get())
        if '$NomeUsina' in p.text:
            p.text = p.text.replace('$NomeUsina', NomeUsina.get())
        if '$TitularUc' in p.text:
            p.text = p.text.replace('$TitularUc', TitularUc.get())
        if '$CNPJ' in p.text:
            p.text = p.text.replace('$CNPJ', CNPJ.get())
        if '$RuaEnel' in p.text:
            p.text = p.text.replace('$RuaEnel', RuaEnel.get())  
        if '$NrEnel' in p.text:
            p.text = p.text.replace('$NrEnel', NrEnel.get())
        if '$BairroEnel' in p.text:
            p.text = p.text.replace('$BairroEnel', BairroEnel.get())
        if '$CepEnel' in p.text:
            p.text = p.text.replace('$CepEnel', CepEnel.get())
        if '$LocalizacaoEnel' in p.text:
            p.text = p.text.replace('$LocalizacaoEnel', LocalizacaoEnel.get())
        if '$ModeloModulo' in p.text:
            p.text = p.text.replace('$ModeloModulo', ModeloModulo.get())
        if '$PotenciaModulo' in p.text:
            p.text = p.text.replace('$PotenciaModulo', PotenciaModulo.get())
        if '$QuantidadeTotalModulos' in p.text:
            p.text = p.text.replace('$QuantidadeTotalModulos', QuantidadeTotalModulos.get())
        if '$ModeloInversor' in p.text:
            p.text = p.text.replace('$ModeloInversor', ModeloInversor.get())
        if '$PotenciaInversor' in p.text:
            p.text = p.text.replace('$PotenciaInversor', PotenciaInversor.get())   
        if '$QuantidadeInversor' in p.text:
            p.text = p.text.replace('$QuantidadeInversor', QuantidadeInversor.get())
        if '$TipoDeEstrutura' in p.text:
            p.text = p.text.replace('$TipoDeEstrutura', TipoDeEstrutura.get())
        if '$AreaTotal' in p.text:
            p.text = p.text.replace('$AreaTotal', AreaTotal)
        if '$FabricanteModulo' in p.text:
            p.text = p.text.replace('$FabricanteModulo', FabricanteModulo.get())
        if '$QuantidadeKwp' in p.text:
            p.text = p.text.replace('$QuantidadeKwp', QuantidadeKwp)
        if '$QuantidadeKwca' in p.text:
            p.text = p.text.replace('$QuantidadeKwca', QuantidadeKwca)
        if '$QuantidadeModulosPorInversor' in p.text:
            p.text = p.text.replace('$QuantidadeModulosPorInversor', QuantidadeModulosPorInversor)
        if '$PotenciaCCPorInversor' in p.text:
            p.text = p.text.replace('$PotenciaCCPorInversor', PotenciaCCPorInversor)
        if '$FabricanteInversor' in p.text:
            p.text = p.text.replace('$FabricanteInversor', FabricanteInversor.get())
        if '$PotenciaTransformador' in p.text:
            p.text = p.text.replace('$PotenciaTransformador', PotenciaTransformador.get())
        if '$TipoTransformador' in p.text:
            p.text = p.text.replace('$TipoTransformador', TipoTransformador.get())
        if '$TensaoSecundario' in p.text:
            p.text = p.text.replace('$TensaoSecundario', TensaoSecundario.get())
        if '$NumeroUc' in p.text:
            p.text = p.text.replace('$NumeroUc', NumeroUc_final)
        if '$ART' in p.text:
            p.text = p.text.replace('$ART', ART.get())
        if '$NivelIsolacao' in p.text:
            p.text = p.text.replace('$NivelIsolacao', NivelIsolacao.get())
        if '$CorrenteTransformador' in p.text:
            p.text = p.text.replace('$CorrenteTransformador', CorrenteCalculadaTransformador)  
        if '$CorrenteSaidaInversor' in p.text:
            p.text = p.text.replace('$CorrenteSaidaInversor', CorrenteSaidaInversor.get())
            
        if QuantidadeTransformadores.get() != '1' and QuantidadeTransformadores.get() !='um':
            if 'O transformador será instalado' in p.text:
                p.text = p.text.replace('O transformador será instalado', 'Os transformadores serão instalados')
            if 'O transformador será alimentado' in p.text:
                p.text = p.text.replace('O transformador será alimentado', 'Os transformadores serão alimentados')
            if 'ao transformador será instalado um painel de média' in p.text:
                p.text = p.text.replace('ao transformador será instalado um painel de média', 'aos transformadores serão instalados painéis de média')
                

    paragraph = document.paragraphs[2]
    paragraph.style = 'titulo'
    paragraph = document.paragraphs[0]
    paragraph.style = 'titulo'
    paragraph = document.paragraphs[1]
    paragraph.style = 'titulo'
    paragraph = document.paragraphs[3]
    paragraph.style = 'titulo'
    
    document.save(output + "\PRJ_" + NumProjeto.get() + "_MD_SE.docx")
    
#************************************ PREENCHENDO MD DE FV *****************************************************     
    caminho2 = caminho.replace("\Python_ENEL_MD_SE_Base_16092022.docx", "\Python_ENEL_MD_FV_Base_16092022.docx") 
    document = Document(caminho2)
    for p in document.paragraphs:
        if '$CidadeEnel' in p.text:
            p.text = p.text.replace('$CidadeEnel', CidadeEnel.get())
        if '$NumProjeto' in p.text:
            p.text = p.text.replace('$NumProjeto', NumProjeto.get())
        if '$NomeUsina' in p.text:
            p.text = p.text.replace('$NomeUsina', NomeUsina.get())
        if '$TitularUc' in p.text:
            p.text = p.text.replace('$TitularUc', TitularUc.get())
        if '$CNPJ' in p.text:
            p.text = p.text.replace('$CNPJ', CNPJ.get())
        if '$RuaEnel' in p.text:
            p.text = p.text.replace('$RuaEnel', RuaEnel.get())  
        if '$NrEnel' in p.text:
            p.text = p.text.replace('$NrEnel', NrEnel.get())
        if '$BairroEnel' in p.text:
            p.text = p.text.replace('$BairroEnel', BairroEnel.get())
        if '$CepEnel' in p.text:
            p.text = p.text.replace('$CepEnel', CepEnel.get())
        if '$LocalizacaoEnel' in p.text:
            p.text = p.text.replace('$LocalizacaoEnel', LocalizacaoEnel.get())
        if '$ModeloModulo' in p.text:
            p.text = p.text.replace('$ModeloModulo', ModeloModulo.get())
        if '$PotenciaModulo' in p.text:
            p.text = p.text.replace('$PotenciaModulo', PotenciaModulo.get())
        if '$QuantidadeTotalModulos' in p.text:
            p.text = p.text.replace('$QuantidadeTotalModulos', QuantidadeTotalModulos.get())
        if '$ModeloInversor' in p.text:
            p.text = p.text.replace('$ModeloInversor', ModeloInversor.get())
        if '$PotenciaInversor' in p.text:
            p.text = p.text.replace('$PotenciaInversor', PotenciaInversor.get())   
        if '$QuantidadeInversor' in p.text:
            p.text = p.text.replace('$QuantidadeInversor', QuantidadeInversor.get())
        if '$TipoDeEstrutura' in p.text:
            p.text = p.text.replace('$TipoDeEstrutura', TipoDeEstrutura.get())
        if '$AreaTotal' in p.text:
            p.text = p.text.replace('$AreaTotal', AreaTotal)
        if '$FabricanteModulo' in p.text:
            p.text = p.text.replace('$FabricanteModulo', FabricanteModulo.get())
        if '$QuantidadeKwp' in p.text:
            p.text = p.text.replace('$QuantidadeKwp', QuantidadeKwp)
        if '$QuantidadeKwca' in p.text:
            p.text = p.text.replace('$QuantidadeKwca', QuantidadeKwca)
        if '$QuantidadeModulosPorInversor' in p.text:
            p.text = p.text.replace('$QuantidadeModulosPorInversor', QuantidadeModulosPorInversor)
        if '$PotenciaCCPorInversor' in p.text:
            p.text = p.text.replace('$PotenciaCCPorInversor', PotenciaCCPorInversor)
        if '$FabricanteInversor' in p.text:
            p.text = p.text.replace('$FabricanteInversor', FabricanteInversor.get())
        if '$TensaoAtendimento' in p.text:
            p.text = p.text.replace('$TensaoAtendimento', TensaoAtendimento)
        if '$Taps' in p.text:
            p.text = p.text.replace('$Taps', Taps)
        if '$PotenciaTransformador' in p.text:
            p.text = p.text.replace('$PotenciaTransformador', PotenciaTransformador.get())
        if '$TipoTransformador' in p.text:
            p.text = p.text.replace('$TipoTransformador', TipoTransformador.get())
        if '$TensaoSecundario' in p.text:
            p.text = p.text.replace('$TensaoSecundario', TensaoSecundario.get())
        if '$NumeroUc' in p.text:
            p.text = p.text.replace('$NumeroUc', NumeroUc_final)
        if '$ART' in p.text:
            p.text = p.text.replace('$ART', ART.get())
        if '$NivelIsolacao' in p.text:
            p.text = p.text.replace('$NivelIsolacao', NivelIsolacao.get())
        if '$CorrenteTransformador' in p.text:
            p.text = p.text.replace('$CorrenteTransformador', CorrenteCalculadaTransformador)  
        if '$CorrenteSaidaInversor' in p.text:
            p.text = p.text.replace('$CorrenteSaidaInversor', CorrenteSaidaInversor.get())
            
        if QuantidadeTransformadores.get() != '1' and QuantidadeTransformadores.get() !='um':
            if 'O transformador será instalado' in p.text:
                p.text = p.text.replace('O transformador será instalado', 'Os transformadores serão instalados')
            if 'O transformador será alimentado' in p.text:
                p.text = p.text.replace('O transformador será alimentado', 'Os transformadores serão alimentados')
            if 'ao transformador será instalado um painel de média' in p.text:
                p.text = p.text.replace('ao transformador será instalado um painel de média', 'aos transformadores serão instalados painéis de média')
    styles = document.styles
    style = styles.add_style('titulo', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Arial'
    font.size = Pt(18)
    font.bold = True
    font.color.rgb = RGBColor(19, 40,105)
    paragraph = document.paragraphs[2]
    paragraph.style = 'titulo'
    paragraph = document.paragraphs[0]
    paragraph.style = 'titulo'
    paragraph = document.paragraphs[1]
    paragraph.style = 'titulo'
    paragraph = document.paragraphs[3]
    paragraph.style = 'titulo'
    document.save(output + "\PRJ_" + NumProjeto.get() + "_MD_FV.docx")
#************************************ PREENCHENDO CRONOGRAMA *****************************************************     
    if len(PrevisaoEnergizacao.get()) == 0 :
        locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')
        caminho8 = caminho.replace("\Python_ENEL_MD_SE_Base_16092022.docx", "\Python_Cronograma_Base_06092022.docx") 
        document = Document(caminho8)
   
        InicioParecer = date.today()
        
        if int(InicioParecer.strftime("%d")) > 5:
            
            InicioParecer = (InicioParecer.replace(day=1) + datetime.timedelta(days=32)).replace(day=1)

        
        FinalParecer = InicioParecer + timedelta(60)
        Diasparecer = (FinalParecer - InicioParecer).days

        InicioMobilizacao = FinalParecer
        FinalMobilizacao = FinalParecer + timedelta(15)
        DiasMob = (FinalMobilizacao - InicioMobilizacao).days

        InicioMontagem = FinalParecer
        FinalMontagem = FinalParecer + timedelta(15)
        DiasMontagem = (FinalMontagem - InicioMontagem).days

        InicioInstalacaoEstruturas = FinalParecer + timedelta(7)
        FinalInstalacaoEstruturas = InicioInstalacaoEstruturas + timedelta(21)
        DiasInstalacaoEstruturas = (FinalInstalacaoEstruturas - InicioInstalacaoEstruturas).days

        InicioInstalacaoMecanica = InicioInstalacaoEstruturas + timedelta(7)
        FinalInstalacaoMecanica = InicioInstalacaoMecanica + timedelta(21)
        DiasInstalacaoMec = (FinalInstalacaoMecanica - InicioInstalacaoMecanica).days

        InicioInstalacaoEletrica = InicioInstalacaoMecanica  + timedelta(7)
        FinalInstalacaoEletrica =  InicioInstalacaoEletrica + timedelta(22)
        DiasInstalacaoEle = (FinalInstalacaoEletrica - InicioInstalacaoEletrica).days

        InicioAdequacao = InicioInstalacaoEletrica + timedelta(9)
        FinalAdequacao =  InicioAdequacao + timedelta(22)
        DiasAdequa = (FinalAdequacao - InicioAdequacao).days

        InicioComissionamento = InicioAdequacao + timedelta(14)
        FinalComissionamento = InicioComissionamento + timedelta(14)
        DiasComissionamento = (FinalComissionamento - InicioComissionamento ).days

        InicioEnergizacao = FinalComissionamento
        FinalEnergizacao = FinalComissionamento + timedelta(7)
        DiasEnergizacao = ( FinalEnergizacao - InicioEnergizacao  ).days

        mesAno1 =  InicioParecer.strftime("%b/%y")
        mesAno2 =  (FinalParecer - timedelta(5)).strftime("%b/%y")
        mesAno3 =  (FinalParecer + timedelta(5)).strftime("%b/%y")
        mesAno4 =  FinalInstalacaoMecanica.strftime("%b/%y")
        mesAno5 =  FinalEnergizacao.strftime("%b/%y")

        InicioEnergizacao = InicioEnergizacao.strftime("%d/%m/%y")
        FinalEnergizacao = FinalEnergizacao.strftime("%d/%m/%y")
        InicioParecer = InicioParecer.strftime("%d/%m/%y")
        FinalParecer = FinalParecer.strftime("%d/%m/%y")
        InicioMobilizacao = InicioMobilizacao.strftime("%d/%m/%y")
        FinalMobilizacao = FinalMobilizacao.strftime("%d/%m/%y")
        InicioMontagem = InicioMontagem.strftime("%d/%m/%y")
        FinalMontagem = FinalMontagem.strftime("%d/%m/%y")
        InicioInstalacaoEstruturas = InicioInstalacaoEstruturas.strftime("%d/%m/%y")
        FinalInstalacaoEstruturas = FinalInstalacaoEstruturas.strftime("%d/%m/%y")
        InicioInstalacaoMecanica = InicioInstalacaoMecanica.strftime("%d/%m/%y")
        FinalInstalacaoMecanica = FinalInstalacaoMecanica.strftime("%d/%m/%y")
        InicioInstalacaoEletrica = InicioInstalacaoEletrica.strftime("%d/%m/%y")
        FinalInstalacaoEletrica = FinalInstalacaoEletrica.strftime("%d/%m/%y")
        InicioAdequacao = InicioAdequacao.strftime("%d/%m/%y")
        FinalAdequacao = FinalAdequacao.strftime("%d/%m/%y")
        InicioComissionamento = InicioComissionamento.strftime("%d/%m/%y")
        FinalComissionamento = FinalComissionamento.strftime("%d/%m/%y")
        for p in document.paragraphs:

            if '$NomeUsina' in p.text:
                p.text = p.text.replace('$NomeUsina', NomeUsina.get())

            if '$TitularUc' in p.text:
                p.text = p.text.replace('$TitularUc', TitularUc.get())
            if '$data' in p.text:
                p.text = p.text.replace('$data', dataAtual)

        for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = cell.text.replace('$mesAno1', mesAno1)
                            cell.text = cell.text.replace('$mesAno2', mesAno2)
                            cell.text = cell.text.replace('$mesAno3', mesAno3)
                            cell.text = cell.text.replace('$mesAno4', mesAno4)
                            cell.text = cell.text.replace('$mesAno5', mesAno5)

                            cell.text = cell.text.replace('$InicioParecer', InicioParecer)
                            cell.text = cell.text.replace('$FinalParecer', FinalParecer)
                            cell.text = cell.text.replace('$InicioMobilizacao', InicioMobilizacao)
                            cell.text = cell.text.replace('$FinalMobilizacao', FinalMobilizacao)
                            cell.text = cell.text.replace('$InicioMontagem',InicioMontagem)
                            cell.text = cell.text.replace('$FinalMontagem', FinalMontagem)
                            cell.text = cell.text.replace('$InicioInstalacaoEstruturas', InicioInstalacaoEstruturas)
                            cell.text = cell.text.replace('$FinalInstalacaoEstruturas', FinalInstalacaoEstruturas)
                            cell.text = cell.text.replace('$InicioInstalacaoMecanica', InicioInstalacaoMecanica)
                            cell.text = cell.text.replace('$FinalInstalacaoMecanica', FinalInstalacaoMecanica)        
                            cell.text = cell.text.replace('$InicioInstalacaoEletrica', InicioInstalacaoEletrica)
                            cell.text = cell.text.replace('$FinalInstalacaoEletrica', FinalInstalacaoEletrica)
                            cell.text = cell.text.replace('$InicioAdequacao', InicioAdequacao)
                            cell.text = cell.text.replace('$FinalAdequacao', FinalAdequacao)
                            cell.text = cell.text.replace('$InicioComissionamento', InicioComissionamento)
                            cell.text = cell.text.replace('$FinalComissionamento', FinalComissionamento)
                            cell.text = cell.text.replace('$InicioEnergizacao', InicioEnergizacao)
                            cell.text = cell.text.replace('$FinalEnergizacao', FinalEnergizacao)

                            cell.text = cell.text.replace('$Diasparecer', str(Diasparecer))
                            cell.text = cell.text.replace('$Diasparecer', str(Diasparecer))
                            cell.text = cell.text.replace('$DiasMob', str(DiasMob))
                            cell.text = cell.text.replace('$DiasMontagem', str(DiasMontagem))
                            cell.text = cell.text.replace('$DiasInstalacaoEstruturas', str(DiasInstalacaoEstruturas))
                            cell.text = cell.text.replace('$DiasInstalacaoMec', str(DiasInstalacaoMec))
                            cell.text = cell.text.replace('$DiasInstalacaoEle', str(DiasInstalacaoEle))
                            cell.text = cell.text.replace('$DiasAdequa', str(DiasAdequa))
                            cell.text = cell.text.replace('$DiasComissionamento', str(DiasComissionamento))
                            cell.text = cell.text.replace('$DiasEnergizacao', str(DiasEnergizacao))
        
        document.save(output + "\PRJ_" + NumProjeto.get() + "_Cronograma.docx")    

#************************************ LISTA DE MATERIAIS ************************************   
    
    if len(String_ModeloTP1.get()) == 0:
            String_ModeloTP1.set("BDEC - FF26S")
    if len(String_FabricanteTP1.get()) == 0:
            String_FabricanteTP1.set("ISOLET")
    if len(String_ModeloTP2.get()) == 0:
            String_ModeloTP2.set("BDE - FT26C")
    planilha = caminhoIcone.replace("\Python_ENEL_MD_SE_Base_16092022.docx", "\Python_ENEL_Anexo_C_Base_16092022.xlsx") 
        
    wb = load_workbook(filename = planilha)

    sh = wb['Plan1']

    for row in sh.iter_cols():
            for i in row:
                b = str(i.value)
                if '$NomeUsina' in b :
                    b = b.replace('$NomeUsina' , NomeUsina.get())
                    i.value = b
                if '$TitularUc' in b :
                    b = b.replace('$TitularUc' , TitularUc.get())
                    i.value = b
                if '$data' in b :
                    b = b.replace('$data' , dataAtual)
                    i.value = b
                if '$ModeloTP1' in b :
                    b = b.replace('$ModeloTP1' , String_ModeloTP1.get())
                    i.value = b
                if '$FabricanteTP1' in b :
                    b = b.replace('$FabricanteTP1' , String_FabricanteTP1.get())
                    i.value = b
                if '$ModeloTP2' in b :
                    b = b.replace('$ModeloTP2' , String_ModeloTP2.get())
                    i.value = b
                if '$FabricanteRele' in b :
                    b = b.replace('$FabricanteRele' , FabricanteRele.get())
                    i.value = b
                if '$Rele' in b :
                    b = b.replace('$Rele' , Rele.get())
                    i.value = b
                if '$TipoTransformador' in b :
                    b = b.replace('$TipoTransformador' , TipoTransformador.get())
                    i.value = b
                if '$PotenciaTransformador' in b :
                    b = b.replace('$PotenciaTransformador' , PotenciaTransformador.get())
                    i.value = b
                if '$TensaoSecundario' in b :
                    b = b.replace('$TensaoSecundario' , TensaoSecundario.get())
                    i.value = b
                if '$DisjuntorGeral' in b :
                    b = b.replace('$DisjuntorGeral' , DisjuntorGeral.get())
                    i.value = b
                if '$DisjuntorParcial' in b :
                    b = b.replace('$DisjuntorParcial' , DisjuntorParcial.get())
                    i.value = b
                if '$QtdDisjuntoresParciais' in b :
                    b = b.replace('$QtdDisjuntoresParciais' , QtdDisjuntoresParciais.get())
                    i.value = b
                if '$FabricanteTC' in b :
                    b = b.replace('$FabricanteTC' , FabricanteTC.get())
                    i.value = b  
                if '$TC' in b :
                    b = b.replace('$TC' , TC.get())
                    i.value = b 

    wb.save(output + "\PRJ_" + NumProjeto.get() + "_LISTA-MATERIAIS_36.2kV.xlsx")    
        
         
    '''if pdf.get() == 1:
        convert(output + "\PRJ_" + NumProjeto.get() + "_Cronograma.docx" , output + "\PRJ_" + NumProjeto.get() + "_Cronograma.pdf")
        convert(output + "\PRJ_" + NumProjeto.get() + "_MD.docx" , output + "\PRJ_" + NumProjeto.get() + "_MD.pdf")
        convert(output + "\PRJ_" + NumProjeto.get() + "_CTC.docx" , output + "\PRJ_" + NumProjeto.get() + "_CTC.pdf")
        convert(output + "\PRJ_" + NumProjeto.get() + "_CUSD.docx" , output + "\PRJ_" + NumProjeto.get() + "_CUSD.pdf")
        convert(output + "\PRJ_" + NumProjeto.get() + "_QC.docx" , output + "\PRJ_" + NumProjeto.get() + "_QC.pdf")
        convert(output + "\PRJ_" + NumProjeto.get() + "_CAP.docx" , output + "\PRJ_" + NumProjeto.get() + "_CAP.pdf")
        convert(output + "\PRJ_" + NumProjeto.get() + "_ANEXOIV.docx" , output + "\PRJ_" + NumProjeto.get() + "_ANEXOIV.pdf")
        convert(output + "\PRJ_" + NumProjeto.get() + "_RCG.docx" , output + "\PRJ_" + NumProjeto.get() + "_RCG.pdf") '''
    messagebox.showinfo(message="Executado com sucesso!")  
     
    
    

def Escolher_saida():
    global output
    output =filedialog.askdirectory()
    return output

def EnderecoDinamicoUsina(event):
    global address
    
    try:

        address = get_address_from_cep(CepEnel.get(), webservice=WebService.VIACEP)
        CidadeEnel.delete('0', 'end')
        CidadeEnel.insert(0,address['cidade'])


        RuaEnel.delete('0', 'end')
        RuaEnel.insert(0,address['logradouro'])


        BairroEnel.delete('0', 'end')
        BairroEnel.insert(0,address['bairro'])

    except exceptions.InvalidCEP as eic:
        messagebox.showinfo(message="CEP inválido.")

    except exceptions.CEPNotFound as ecnf:
        messagebox.showinfo(message="CEP não encontrado 😕")

    except exceptions.ConnectionError as errc:
        messagebox.showinfo(message="Não foi possível preencher automaticamente.\n\n Sem conexão com a internet 🌐❌")

    except exceptions.Timeout as errt:
        messagebox.showinfo(message="CEP não encontrado 😕")

    except exceptions.HTTPError as errh:
        messagebox.showinfo(message="CEP não encontrado 😕")

    except exceptions.BaseException as e:
        messagebox.showinfo(message="CEP não encontrado 😕")
    
    
def EnderecoDinamicoTitular(event):
    global address
    
    try:

        addressTitular = get_address_from_cep(CepTitular.get(), webservice=WebService.VIACEP)
        CidadeTitular.delete('0', 'end')
        CidadeTitular.insert(0,addressTitular['cidade'])


        RuaTitular.delete('0', 'end')
        RuaTitular.insert(0,addressTitular['logradouro'])


        BairroTitular.delete('0', 'end')
        BairroTitular.insert(0,addressTitular['bairro'])
        
        UfTitular.delete('0', 'end')
        UfTitular.insert(0,addressTitular['uf'])

    except exceptions.InvalidCEP as eic:
        messagebox.showinfo(message="CEP inválido.")

    except exceptions.CEPNotFound as ecnf:
        messagebox.showinfo(message="CEP não encontrado 😕")

    except exceptions.ConnectionError as errc:
        messagebox.showinfo(message="Não foi possível preencher automaticamente.\n\n Sem conexão com a internet 🌐❌")

    except exceptions.Timeout as errt:
        messagebox.showinfo(message="CEP não encontrado 😕")

    except exceptions.HTTPError as errh:
        messagebox.showinfo(message="CEP não encontrado 😕")

    except exceptions.BaseException as e:
        messagebox.showinfo(message="CEP não encontrado 😕")





#INSERINDO DADOS PARA A CONSTRUÇÃO DO MEMORIAL -----------------------------------------------------------------------
#Definindo variáveis
NumProjeto=tk.StringVar() #feito
NomeUsina=tk.StringVar() #feito
CidadeEnel = tk.StringVar() #feito
TitularUc = tk.StringVar()#feito
CNPJ = tk.StringVar()#feito
RuaEnel = tk.StringVar() #feito
NrEnel = tk.StringVar() #feito
BairroEnel = tk.StringVar() #feito
CepEnel = tk.StringVar() #feito
LocalizacaoEnel = tk.StringVar() #feito
ModeloModulo = tk.StringVar() #feito
PotenciaModulo = tk.StringVar() #feito
QuantidadeTotalModulos = tk.StringVar() #feito
ModeloInversor = tk.StringVar() #feito
PotenciaInversor = tk.StringVar() #feito
QuantidadeInversor = tk.StringVar() #feito
TipoDeEstrutura = tk.StringVar() #feito

FabricanteModulo = tk.StringVar()  #feito
FabricanteInversor = tk.StringVar() #feito
PotenciaTransformador = tk.StringVar() #feito
TipoTransformador = tk.StringVar() #feito
TensaoSecundario = tk.StringVar() #feito
NumeroUc = tk.StringVar()
Taps = tk.StringVar() 
Uc = tk.IntVar() 
pdf = tk.IntVar() 
AlturaModulo =tk.StringVar()
LarguraModulo = tk.StringVar()


#INSERINDO DADOS PARA A CONSTRUÇÃO DO CTC -----------------------------------------------------------------------
#Definindo variáveis
ART = tk.StringVar()#feito
NrTitular = tk.StringVar() #feito
BairroTitular = tk.StringVar()#feito
RuaTitular = tk.StringVar()#feito
CidadeTitular = tk.StringVar() #feito
CepTitular = tk.StringVar()#feito
TelefoneTitular = tk.StringVar() #feito
EmailTitular = tk.StringVar()#feito 
DescricaoAtividade = tk.StringVar() #feito 
CodigoCNAE = tk.StringVar() #feito 
Rele = tk.StringVar() #feito
FabricanteRele = tk.StringVar() #feito
TC = tk.StringVar() #feito
FabricanteTC = tk.StringVar()#feito 
NivelIsolacao = tk.StringVar() #feito
CorrenteSaidaInversor = tk.StringVar()#feito  
QuantidadeTransformadores = tk.StringVar() #feito  
PotenciaTotalTransformadores = tk.StringVar()    

Impedancia = tk.StringVar() #feito 

#INSERINDO DADOS PARA A CONSTRUÇÃO DO CUSD -----------------------------------------------------------------------
#Definindo variáveis
ComplementoEnel = tk.StringVar() #feito 
InscricaoEstadual = tk.StringVar() #feito 
InscricaoMunicipal = tk.StringVar() #feito 
ComplementoTitular = tk.StringVar() #feito 
UfTitular = tk.StringVar() #feito 

RepresentanteLegalUm = tk.StringVar() #feito 
CPFRepresentanteLegalUm = tk.StringVar() #feito 
RGRepresentanteLegalUm = tk.StringVar() #feito 
OrgaoEmissorRGUm = tk.StringVar() #feito 
CargoUm = tk.StringVar() #feito 
CotaUm = tk.StringVar() #feito 
TelefoneCelularRepresentanteLegalUm = tk.StringVar() #feito 
EmailRepresentanteLegalUm = tk.StringVar() #feito 

RepresentanteLegalDois = tk.StringVar() #feito 
CPFRepresentanteLegalDois = tk.StringVar() #feito 
RGRepresentanteLegalDois = tk.StringVar() #feito 
OrgaoEmissorRGDois = tk.StringVar() #feito 
CargoDois = tk.StringVar() #feito 
CotaDois = tk.StringVar() #feito 
TelefoneCelularRepresentanteLegalDois = tk.StringVar() #feito 
EmailRepresentanteLegalDois = tk.StringVar() #feito 
#INSERINDO DADOS PARA A CONSTRUÇÃO DO CAP -----------------------------------------------------------------------
PrevisaoEnergizacao = tk.StringVar() #feito 

DisjuntorGeral = tk.StringVar()#feito
DisjuntorParcial = tk.StringVar()#feito
QtdDisjuntoresParciais = tk.StringVar()#feito
String_ModeloTP1 = tk.StringVar()#feito
String_FabricanteTP1 = tk.StringVar()#feito
String_ModeloTP2 = tk.StringVar()#feito







menubar = Menu(root)
filemenu = Menu(menubar, tearoff=0)
filemenu.add_command(label="Escolher local de saída", command=Escolher_saida)
filemenu.add_command(label="Gerar Arquivos", command=imp)
filemenu.add_command(label="Função Backup", command=backup)
menubar.add_cascade(label="Gerar Arquivos", menu=filemenu)
#Inserindo Marcador
myLabel=Label(second_frame,text="Dados Gerais 💬 ", font=('Arial', 13,'bold'),pady=20, fg = 'orange' , bg='white')
myLabel.grid(row=3,column=1, sticky='w')
#Inserindo Marcador
myLabel=Label(second_frame,text="Configuração da Usina ⚙️ ", font=('Arial', 13,'bold'),pady=20, fg = 'blue' , bg='white')
myLabel.grid(row=3,column=3, sticky='w')
#Inserindo Marcador
myLabel=Label(second_frame,text="Localização da Usina 🏳️ ", pady = 20,font=('Arial', 13,'bold'), fg = 'green' , bg='white')
myLabel.grid(row=10,column=1, sticky='w')
#Inserindo Marcador
myLabel=Label(second_frame,text="Dados do Titular 👤 ", pady = 20,font=('Arial', 13,'bold'), fg = 'purple' , bg='white')
myLabel.grid(row=18,column=1, sticky='w')
#Inserindo Marcador
myLabel=Label(second_frame,text=" Representante Legal 1 📝 ", font=('Arial', 13,'bold'),pady=20, fg = 'red' , bg='white')
myLabel.grid(row=36,column=1, sticky='w')
#Inserindo Marcador
myLabel=Label(second_frame,text=" Representante Legal 2 📝 ", font=('Arial', 13,'bold'),pady=20, fg = 'red' , bg='white')
myLabel.grid(row=36,column=3, sticky='sw')

#Inserindo Marcador
myLabel=Label(second_frame,pady=10, bg='white')
myLabel.grid(row=0,column=1)


'''****************************************DADOS GERAIS **************************************************'''

#------------------------------------- Numero prj
NumeroProjeto_label = tk.Label(second_frame, 
         text="Número do projeto :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=4,sticky='e')
NumeroProjeto= tk.Entry(second_frame, textvariable =  NumProjeto ,width = 15,font=('Arial 10'),borderwidth=5)

NumeroProjeto.grid(row=4, column=1,sticky='w')

#------------------------------------- Nome usina
NomeUsina_label = tk.Label(second_frame, 
         text="Nome da usina :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=5,sticky='e')
NomeUsina= tk.Entry(second_frame, textvariable =  NomeUsina ,width = 30,font=('Arial 10'),borderwidth=5)
NomeUsina .grid(row=5, column=1,sticky='w')

#------------------------------------- ART
ART_entry = tk.Label(second_frame, 
         text="ART :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=6,sticky='e')
ART = tk.Entry(second_frame, textvariable =  ART ,width = 30,font=('Arial 10'),borderwidth=5)
ART.grid(row=6, column=1,sticky='w') 

#------------------------------------- NumeroUc
NumeroUc_entry = tk.Label(second_frame, 
         text="Número da UC :",font=('Arial', 10,'bold'),pady=5,padx = 10 ,fg="black", bg="white",
                borderwidth=0).grid(row=7,sticky='e')
NumeroUc = tk.Entry(second_frame, textvariable = NumeroUc ,width = 12,font=('Arial 10'),borderwidth=5)
NumeroUc.grid(row=7, column=1,sticky='w') 

Uc.set(-1)

R2 = Radiobutton(second_frame, text="Uc não existente", variable=Uc, value=1,padx=-20,fg="black", bg="white")                 
R2.grid(row=7, column=1,sticky='e')


#------------------------------------- CNPJ
CNPJ_entry = tk.Label(second_frame, 
         text="CNPJ :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=8,sticky='e')
CNPJ = tk.Entry(second_frame, textvariable = CNPJ ,width = 30,font=('Arial 10'),borderwidth=5)
CNPJ.grid(row=8, column=1,sticky='w') 




'''****************************************LOCALIZAÇÃO DA USINA **************************************************'''
#------------------------------------- RUA
RuaEnel_entry = tk.Label(second_frame, 
         text="Rua :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=11,sticky='e')
RuaEnel = tk.Entry(second_frame, textvariable = RuaEnel ,width = 30,font=('Arial 10'),borderwidth=5)
RuaEnel.bind("<FocusIn>", lambda args: RuaEnel.delete('0', 'end'))
RuaEnel.grid(row=11, column=1,sticky='w') 
#------------------------------------- Cidade
CidadeEnel_entry = tk.Label(second_frame, 
         text="Cidade :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=12,sticky='e')
CidadeEnel = tk.Entry(second_frame, textvariable =  CidadeEnel ,width = 30,font=('Arial 10'),borderwidth=5)
CidadeEnel.bind("<FocusIn>", lambda args: CidadeEnel.delete('0', 'end'))
CidadeEnel.grid(row=12, column=1,sticky='w') 

#------------------------------------- Número
NrEnel_entry = tk.Label(second_frame, 
         text="N° :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=13,sticky='e')
NrEnel = tk.Entry(second_frame, textvariable =  NrEnel ,width = 10,font=('Arial 10'),borderwidth=5)

NrEnel.grid(row=13, column=1,sticky='w')

#------------------------------------- Bairro
BairroEnel_entry = tk.Label(second_frame, 
         text="Bairro :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=14,sticky='e')
BairroEnel = tk.Entry(second_frame, textvariable =  BairroEnel ,width = 30,font=('Arial 10'),borderwidth=5)
BairroEnel.bind("<FocusIn>", lambda args: BairroEnel.delete('0', 'end'))
BairroEnel.grid(row=14, column=1,sticky='w')
#------------------------------------- CEP
CepEnel_entry = tk.Label(second_frame, 
         text="CEP :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=15,sticky='e')
CepEnel = tk.Entry(second_frame, textvariable =  CepEnel ,width = 18,font=('Arial 10'),borderwidth=5)
CepEnel.bind("<FocusOut>",EnderecoDinamicoUsina)
CepEnel.grid(row=15, column=1,sticky='w')

#------------------------------------- Complemento
ComplementoEnel_entry = tk.Label(second_frame, 
         text="Complemento :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=16,sticky='e')
ComplementoEnel = tk.Entry(second_frame, textvariable =  ComplementoEnel ,width = 30,font=('Arial 10'),borderwidth=5)
ComplementoEnel.grid(row=16, column=1,sticky='w')

#------------------------------------- LocalizacaoEnel
LocalizacaoEnel_entry = tk.Label(second_frame, 
         text="Coordenadas :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=17,sticky='e')
LocalizacaoEnel = tk.Entry(second_frame, textvariable =  LocalizacaoEnel ,width = 30,font=('Arial 10'),borderwidth=5)
LocalizacaoEnel.grid(row=17, column=1,sticky='w')


'''****************************************CONFIGURAÇÃO DA USINA **************************************************'''

#------------------------------------- ModeloModulo
ModeloModulo_entry = tk.Label(second_frame, 
         text="Modelo do módulo :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=4,column =2,sticky='e')
ModeloModulo = tk.Entry(second_frame, textvariable =  ModeloModulo ,width = 30,font=('Arial 10'),borderwidth=5)
ModeloModulo.grid(row=4, column=3,sticky='w')

#------------------------------------- PotenciaModulo
PotenciaModulo_entry = tk.Label(second_frame, 
         text="Potência do módulo (W):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=5,column =2,sticky='e')
PotenciaModulo = tk.Entry(second_frame, textvariable =  PotenciaModulo ,width = 10,font=('Arial 10'),borderwidth=5)
PotenciaModulo.grid(row=5, column=3,sticky='w')

#------------------------------------- QuantidadeTotalModulos
QuantidadeTotalModulos_entry = tk.Label(second_frame, 
         text="Quantidade total de módulos :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=6,column =2,sticky='e')
QuantidadeTotalModulos = tk.Entry(second_frame, textvariable =  QuantidadeTotalModulos ,width = 15,font=('Arial 10'),borderwidth=5)
QuantidadeTotalModulos.grid(row=6, column=3,sticky='w')

#------------------------------------- FabricanteModulo
FabricanteModulo_entry = tk.Label(second_frame, 
         text="Fabricante do módulo :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=7,column =2,sticky='e')
FabricanteModulo = tk.Entry(second_frame, textvariable =  FabricanteModulo ,width = 20,font=('Arial 10'),borderwidth=5)
FabricanteModulo.grid(row=7, column=3,sticky='w')


#------------------------------------- PotenciaInversor
PotenciaInversor_entry = tk.Label(second_frame, 
         text="Potência do inversor (kVA):",font=('Arial', 10,'bold'),padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=8,column =2,sticky='e')
PotenciaInversor = tk.Entry(second_frame, textvariable =  PotenciaInversor ,width = 10,font=('Arial 10'),borderwidth=5)
PotenciaInversor.grid(row=8, column=3,sticky='w')

#------------------------------------- QuantidadeInversor
QuantidadeInversor_entry = tk.Label(second_frame, 
         text="Quantidade de inversores :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=9,column =2,sticky='e')
QuantidadeInversor = tk.Entry(second_frame, textvariable =  QuantidadeInversor ,width = 6,font=('Arial 10'),borderwidth=5)
QuantidadeInversor.grid(row=9, column=3,sticky='w')

#------------------------------------- ModeloInversor
ModeloInversor_entry = tk.Label(second_frame, 
         text="Modelo do inversor :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=10,column =2,sticky='e')
ModeloInversor = tk.Entry(second_frame, textvariable =  ModeloInversor ,width = 30,font=('Arial 10'),borderwidth=5)
ModeloInversor.grid(row=10, column=3,sticky='w')

#------------------------------------- FabricanteInversor
FabricanteInversor_entry = tk.Label(second_frame, 
         text="Fabricante do inversor :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=11,column =2,sticky='e')
FabricanteInversor = tk.Entry(second_frame, textvariable =  FabricanteInversor ,width = 20,font=('Arial 10'),borderwidth=5)
FabricanteInversor.grid(row=11, column=3,sticky='w')

#-----------------------------------CorrenteSaidaInversor
CorrenteSaidaInversor_entry = tk.Label(second_frame, 
         text="Corrente de saída do inversor (A):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=12,column =2,sticky='e')
CorrenteSaidaInversor = tk.Entry(second_frame, textvariable = CorrenteSaidaInversor ,width = 10,font=('Arial 10'),borderwidth=5)
CorrenteSaidaInversor.grid(row=12, column=3,sticky='w')

#------------------------------------- TipoDeEstrutura
TipoDeEstrutura_entry = tk.Label(second_frame, 
         text="Tipo de estrutura :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=13,column =2,sticky='e')
TipoDeEstrutura = tk.Entry(second_frame, textvariable =  TipoDeEstrutura ,width = 20,font=('Arial 10'),borderwidth=5)
TipoDeEstrutura.grid(row=13, column=3,sticky='w')
#------------------------------------- AreaTotal
AreaTotal_entry = tk.Label(second_frame, 
         text="Área total coberta por módulos (m²) :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=14,column =2,sticky='e')
LarguraModulo_entry = tk.Label(second_frame, 
         text="Largura (m) ",font=('Arial', 10,'bold'),pady=5,fg="black", bg="white",
                borderwidth=0).grid(row=14,column =3,sticky='w')
LarguraModulo = tk.Entry(second_frame, textvariable = LarguraModulo ,width = 7,font=('Arial 10'),borderwidth=5)
LarguraModulo.grid(row=14, column=3)

AlturaModulo_entry = tk.Label(second_frame, 
         text="Altura (m) ",font=('Arial', 10,'bold'),pady=5,fg="black", bg="white",
                borderwidth=0).grid(row=14,column =3,sticky='e')
AlturaModulo = tk.Entry(second_frame, textvariable = AlturaModulo ,width = 7,font=('Arial 10'),borderwidth=5)
AlturaModulo.grid(row=14, column=4)


#-----------------------------------PotenciaTransformador
PotenciaTransformador_entry = tk.Label(second_frame, 
         text="Potência do transformador (kVA) :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=15,column =2,sticky='e')
PotenciaTransformador = tk.Entry(second_frame, textvariable =  PotenciaTransformador ,width = 15,font=('Arial 10'),borderwidth=5)
PotenciaTransformador.grid(row=15, column=3,sticky='w')

#-----------------------------------TipoTransformador
TipoTransformador_entry = tk.Label(second_frame, 
         text="Tipo do transformador :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=16,column =2,sticky='e')
TipoTransformador = tk.Entry(second_frame, textvariable = TipoTransformador ,width = 10,font=('Arial 10'),borderwidth=5)
TipoTransformador.grid(row=16, column=3,sticky='w')

#-----------------------------------TensaoSecundario
TensaoSecundario_entry = tk.Label(second_frame, 
         text="Tensão no secundário do transf. (V) :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=17,column =2,sticky='e')
TensaoSecundario = tk.Entry(second_frame, textvariable = TensaoSecundario ,width = 10,font=('Arial 10'),borderwidth=5)
TensaoSecundario.grid(row=17, column=3,sticky='w')

#-----------------------------------QuantidadeTransformadores
QuantidadeTransformadores_entry = tk.Label(second_frame, 
         text="Quantidade de transformadores :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=18,column =2,sticky='e')
QuantidadeTransformadores = tk.Entry(second_frame, textvariable = QuantidadeTransformadores ,width = 6,font=('Arial 10'),borderwidth=5)
QuantidadeTransformadores.grid(row=18, column=3,sticky='w')



#-----------------------------------Impedancia 
Impedancia_entry = tk.Label(second_frame, 
         text="Impedância do transformador(%) :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=19,column =2,sticky='e')
Impedancia  = tk.Entry(second_frame, textvariable = Impedancia  ,width = 10,font=('Arial 10'),borderwidth=5)
Impedancia.grid(row=19, column=3,sticky='w')

#-----------------------------------NivelIsolacao 
NivelIsolacao_entry = tk.Label(second_frame, 
         text="Nível de isolação dos condutores (kV):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=20,column =2,sticky='e')
NivelIsolacao  = tk.Entry(second_frame, textvariable = NivelIsolacao  ,width = 10,font=('Arial 10'),borderwidth=5)
NivelIsolacao.grid(row=20, column=3,sticky='w')
#-----------------------------------Rele 
Rele_entry = tk.Label(second_frame, 
         text="Modelo do relé :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=21,column =2,sticky='e')
Rele = tk.Entry(second_frame, textvariable = Rele ,width = 20,font=('Arial 10'),borderwidth=5)
Rele.grid(row=21, column=3,sticky='w')
#-----------------------------------FabricanteRele
FabricanteRele_entry = tk.Label(second_frame, 
         text="Fabricante do relé :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=22,column =2,sticky='e')
FabricanteRele = tk.Entry(second_frame, textvariable = FabricanteRele ,width = 20,font=('Arial 10'),borderwidth=5)
FabricanteRele.grid(row=22, column=3,sticky='w')
#-----------------------------------TC
TC_entry = tk.Label(second_frame, 
         text="Modelo do TC :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=23,column =2,sticky='e')
TC = tk.Entry(second_frame, textvariable = TC ,width = 20,font=('Arial 10'),borderwidth=5)
TC.grid(row=23, column=3,sticky='w')
#-----------------------------------FabricanteTC
FabricanteTC_entry = tk.Label(second_frame, 
         text="Fabricante do TC :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=24,column =2,sticky='e')
FabricanteTC = tk.Entry(second_frame, textvariable = FabricanteTC ,width = 20,font=('Arial 10'),borderwidth=5)
FabricanteTC.grid(row=24, column=3,sticky='w')

#-----------------------------------DisjuntorGeral
DisjuntorGeral_entry = tk.Label(second_frame, 
         text="Disjuntor geral do QBGT (A):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=25,column =2,sticky='e')
DisjuntorGeral = tk.Entry(second_frame, textvariable = DisjuntorGeral ,width = 15,font=('Arial 10'),borderwidth=5)
DisjuntorGeral.grid(row=25, column=3,sticky='w')
#-----------------------------------DisjuntorParcial
DisjuntorParcial_entry = tk.Label(second_frame, 
         text="Disjuntor parcial (A):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=26,column =2,sticky='e')
DisjuntorParcial = tk.Entry(second_frame, textvariable = DisjuntorParcial ,width = 15,font=('Arial 10'),borderwidth=5)
DisjuntorParcial.grid(row=26, column=3,sticky='w')

#-----------------------------------QtdDisjuntoresParciais
QtdDisjuntoresParciais_entry = tk.Label(second_frame, 
         text="Quantidade de disjuntores parciais:",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=27,column =2,sticky='e')
QtdDisjuntoresParciais = tk.Entry(second_frame, textvariable = QtdDisjuntoresParciais ,width = 15,font=('Arial 10'),borderwidth=5)
QtdDisjuntoresParciais.grid(row=27, column=3,sticky='w')

#-----------------------------------ModeloTP1
ModeloTP1_entry = tk.Label(second_frame, 
         text="Modelo TP1 (1000VA):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=28,column =2,sticky='e')
ModeloTP1 = tk.Entry(second_frame, textvariable = String_ModeloTP1 ,width = 20,font=('Arial 10'),borderwidth=5)
ModeloTP1.grid(row=28, column=3,sticky='w')

#-----------------------------------FabricanteTP1
FabricanteTP1_entry = tk.Label(second_frame, 
         text="Fabricante TP1 (1000VA):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=29,column =2,sticky='e')
FabricanteTP1 = tk.Entry(second_frame, textvariable = String_FabricanteTP1 ,width = 20,font=('Arial 10'),borderwidth=5)
FabricanteTP1.grid(row=29, column=3,sticky='w')


#-----------------------------------ModeloTP2
ModeloTP2_entry = tk.Label(second_frame, 
         text="Modelo TP2 (500VA):",font=('Arial', 10,'bold'),pady=5,padx=10,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=30,column =2,sticky='e')
ModeloTP2 = tk.Entry(second_frame, textvariable = String_ModeloTP2 ,width = 20,font=('Arial 10'),borderwidth=5)
ModeloTP2.grid(row=30, column=3,sticky='w')

#-----------------------------------Aviso
Aviso_entry = tk.Label(second_frame, 
         text="Caso não seja especificado",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=31,column =2,sticky='e')
Aviso_entry = tk.Label(second_frame, 
         text=" o modelo e fabricante dos TP's :",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=31,column =3,sticky='w')
Aviso_entry = tk.Label(second_frame, 
         text="TP1 25kV: Modelo BDEC - ",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=32,column =2,sticky='e')

Aviso_entry = tk.Label(second_frame, 
         text="FF26S Fabricante ISOLET",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=32,column =3,sticky='w')
Aviso_entry = tk.Label(second_frame, 
         text="TP2 25kV: Modelo BDE - ",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=33,column =2,sticky='e')

Aviso_entry = tk.Label(second_frame, 
         text="FT26C Fabricante ISOLET",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=33,column =3,sticky='w')
Aviso_entry = tk.Label(second_frame, 
         text="TP1 = TP2 13,8kV: Modelo TPB",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=34,column =2,sticky='e')

Aviso_entry = tk.Label(second_frame, 
         text="-15 Fabricante INSTRUMENTI",font=('Arial', 9,'bold'),pady=5,fg="#FF1493", bg="white",
                borderwidth=0).grid(row=34,column =3,sticky='w')



#-----------------------------------PrevisaoEnergizacao
PrevisaoEnergizacao_entry = tk.Label(second_frame, 
         text="Previsão de energização :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=35,column =2,sticky='e')
PrevisaoEnergizacao = tk.Entry(second_frame, textvariable = PrevisaoEnergizacao ,width = 30,font=('Arial 10'),borderwidth=5)
PrevisaoEnergizacao.grid(row=35, column=3,sticky='w')



'''****************************************DADOS DO TITULAR **************************************************'''

#------------------------------------- titular uc
TitularUc_entry = tk.Label(second_frame, 
         text="Titular UC :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=19,sticky='e')
TitularUc = tk.Entry(second_frame, textvariable =  TitularUc ,width = 30,font=('Arial 10'),borderwidth=5)
TitularUc.grid(row=19, column=1,sticky='w') 


#------------------------------------- RuaTitular
RuaTitular_entry = tk.Label(second_frame, 
         text="Rua :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=20,sticky='e')
RuaTitular = tk.Entry(second_frame, textvariable =  RuaTitular ,width = 30,font=('Arial 10'),borderwidth=5)
RuaTitular.bind("<FocusIn>", lambda args: RuaTitular.delete('0', 'end'))
RuaTitular.grid(row=20, column=1,sticky='w') 

#------------------------------------- CidadeTitular
CidadeTitular_entry = tk.Label(second_frame, 
         text="Cidade :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=21,sticky='e')
CidadeTitular = tk.Entry(second_frame, textvariable =  CidadeTitular ,width = 30,font=('Arial 10'),borderwidth=5)
CidadeTitular.bind("<FocusIn>", lambda args: CidadeTitular.delete('0', 'end'))
CidadeTitular.grid(row=21, column=1,sticky='w') 


#------------------------------------- NrTitular
NrTitular_entry = tk.Label(second_frame, 
         text="N° :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=22,sticky='e')
NrTitular = tk.Entry(second_frame, textvariable =  NrTitular ,width = 10,font=('Arial 10'),borderwidth=5)
NrTitular.grid(row=22, column=1,sticky='w') 


#------------------------------------- BairroTitular
BairroTitular_entry = tk.Label(second_frame, 
         text="Bairro :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=23,sticky='e')
BairroTitular = tk.Entry(second_frame, textvariable =  BairroTitular ,width = 30,font=('Arial 10'),borderwidth=5)
BairroTitular.bind("<FocusIn>", lambda args: BairroTitular.delete('0', 'end'))
BairroTitular.grid(row=23, column=1,sticky='w') 

#------------------------------------- CepTitular
CepTitular_entry = tk.Label(second_frame, 
         text="CEP :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=24,sticky='e')
CepTitular = tk.Entry(second_frame, textvariable =  CepTitular ,width = 18,font=('Arial 10'),borderwidth=5)
CepTitular.bind("<FocusOut>",EnderecoDinamicoTitular)
CepTitular.grid(row=24, column=1,sticky='w')

#------------------------------------- ComplementoTitular
ComplementoTitular_entry = tk.Label(second_frame, 
         text="Complemento :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=25,sticky='e')
ComplementoTitular = tk.Entry(second_frame, textvariable = ComplementoTitular ,width = 30,font=('Arial 10'),borderwidth=5)
ComplementoTitular.grid(row=25, column=1,sticky='w')

#------------------------------------- UfTitular
UfTitular_entry = tk.Label(second_frame, 
         text="UF :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=26,sticky='e')
UfTitular = tk.Entry(second_frame, textvariable = UfTitular ,width = 8,font=('Arial 10'),borderwidth=5)
UfTitular.bind("<FocusIn>", lambda args: UfTitular.delete('0', 'end'))
UfTitular.grid(row=26, column=1,sticky='w')


#------------------------------------- TelefoneTitular
TelefoneTitular_entry = tk.Label(second_frame, 
         text="Telefone :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=27,sticky='e')
TelefoneTitular = tk.Entry(second_frame, textvariable =  TelefoneTitular ,width = 25,font=('Arial 10'),borderwidth=5)
TelefoneTitular.grid(row=27, column=1,sticky='w')

#------------------------------------- EmailTitular
EmailTitular_entry = tk.Label(second_frame, 
         text="Email :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=28,sticky='e')
EmailTitular = tk.Entry(second_frame, textvariable =  EmailTitular ,width = 30,font=('Arial 10'),borderwidth=5)
EmailTitular.grid(row=28, column=1,sticky='w')


#------------------------------------- InscricaoMunicipal
InscricaoMunicipal_entry = tk.Label(second_frame, 
         text="Inscrição Municipal :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=29,sticky='e')
InscricaoMunicipal = tk.Entry(second_frame, textvariable = InscricaoMunicipal ,width = 20,font=('Arial 10'),borderwidth=5)
InscricaoMunicipal.grid(row=29, column=1,sticky='w')


#------------------------------------- InscricaoEstadual
InscricaoEstadual_entry = tk.Label(second_frame, 
         text="Inscrição Estadual :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=30,sticky='e')
InscricaoEstadual = tk.Entry(second_frame, textvariable = InscricaoEstadual ,width = 20,font=('Arial 10'),borderwidth=5)
InscricaoEstadual.grid(row=30, column=1,sticky='w')



#------------------------------------- CodigoCNAE
CodigoCNAE_entry = tk.Label(second_frame, 
         text="Código CNAE :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=31,sticky='e')
CodigoCNAE = tk.Entry(second_frame, textvariable =  CodigoCNAE ,width = 30,font=('Arial 10'),borderwidth=5)
CodigoCNAE.grid(row=31, column=1,sticky='w')


#------------------------------------- DescricaoAtividade
DescricaoAtividade_entry = tk.Label(second_frame, 
         text="Descrição da atividade :",font=('Arial', 10,'bold'),pady=2,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=35,sticky='e')
DescricaoAtividade = tk.Text(second_frame,height = 8 ,width=30, font=('Arial 10'),borderwidth=5)
DescricaoAtividade.grid(row=35, column=1,sticky='w')



'''****************************************DADOS DO REPRESENTANTE LEGAL************************************************'''


#-----------------------------------RepresentanteLegalUm
RepresentanteLegalUm_entry = tk.Label(second_frame, 
         text="Nome Completo :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=37,column =0,sticky='e')
RepresentanteLegalUm = tk.Entry(second_frame, textvariable = RepresentanteLegalUm ,width = 30,font=('Arial 10'),borderwidth=5)
RepresentanteLegalUm.grid(row=37, column=1,sticky='w')

#-----------------------------------CPFRepresentanteLegalUm
CPFRepresentanteLegalUm_entry = tk.Label(second_frame, 
         text="CPF :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=38,column =0,sticky='e')
CPFRepresentanteLegalUm = tk.Entry(second_frame, textvariable = CPFRepresentanteLegalUm ,width = 20,font=('Arial 10'),borderwidth=5)
CPFRepresentanteLegalUm.grid(row=38, column=1,sticky='w')


#-----------------------------------RGRepresentanteLegalUm
RGRepresentanteLegalUm_entry = tk.Label(second_frame, 
         text="RG :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=39,column =0,sticky='e')
RGRepresentanteLegalUm = tk.Entry(second_frame, textvariable = RGRepresentanteLegalUm ,width = 20,font=('Arial 10'),borderwidth=5)
RGRepresentanteLegalUm.grid(row=39, column=1,sticky='w')


#-----------------------------------OrgaoEmissorRGUm
OrgaoEmissorRGUm_entry = tk.Label(second_frame, 
         text="Órgão Emissor :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=40,column =0,sticky='e')
OrgaoEmissorRGUm = tk.Entry(second_frame, textvariable = OrgaoEmissorRGUm ,width = 20,font=('Arial 10'),borderwidth=5)
OrgaoEmissorRGUm.grid(row=40, column=1,sticky='w')

#-----------------------------------CargoUm
CargoUm_entry = tk.Label(second_frame, 
         text="Cargo :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=41,column =0,sticky='e')
CargoUm = tk.Entry(second_frame, textvariable = CargoUm ,width = 20,font=('Arial 10'),borderwidth=5)
CargoUm.grid(row=41, column=1,sticky='w')

#-----------------------------------CotaUm
CotaUm = tk.Label(second_frame, 
         text="Cota :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=42,column =0,sticky='e')
CotaUm = tk.Entry(second_frame, textvariable = CotaUm ,width = 20,font=('Arial 10'),borderwidth=5)
CotaUm.grid(row=42, column=1,sticky='w')


#-----------------------------------TelefoneCelularRepresentanteLegalUm
TelefoneCelularRepresentanteLegalUm_entry = tk.Label(second_frame, 
         text="Telefone :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=43,column =0,sticky='e')
TelefoneCelularRepresentanteLegalUm = tk.Entry(second_frame, textvariable = TelefoneCelularRepresentanteLegalUm ,width = 20,font=('Arial 10'),borderwidth=5)
TelefoneCelularRepresentanteLegalUm.grid(row=43, column=1,sticky='w')

#-----------------------------------EmailRepresentanteLegalUm
EmailRepresentanteLegalUm = tk.Label(second_frame, 
         text="Email :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=44,column =0,sticky='e')
EmailRepresentanteLegalUm = tk.Entry(second_frame, textvariable = EmailRepresentanteLegalUm ,width = 20,font=('Arial 10'),borderwidth=5)
EmailRepresentanteLegalUm.grid(row=44, column=1,sticky='w')

#----------------------------------------  Representante legal 2   ---------------------------------

#-----------------------------------RepresentanteLegalDois
RepresentanteLegalDois_entry = tk.Label(second_frame, 
         text="Nome Completo :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=37,column =2,sticky='e')
RepresentanteLegalDois = tk.Entry(second_frame, textvariable = RepresentanteLegalDois ,width = 30,font=('Arial 10'),borderwidth=5)
RepresentanteLegalDois.grid(row=37, column=3,sticky='w')

#-----------------------------------CPFRepresentanteLegalDois
CPFRepresentanteLegalDois_entry = tk.Label(second_frame, 
         text="CPF :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=38,column =2,sticky='e')
CPFRepresentanteLegalDois = tk.Entry(second_frame, textvariable = CPFRepresentanteLegalDois ,width = 20,font=('Arial 10'),borderwidth=5)
CPFRepresentanteLegalDois.grid(row=38, column=3,sticky='w')


#-----------------------------------RGRepresentanteLegalDois
RGRepresentanteLegalDois_entry = tk.Label(second_frame, 
         text="RG :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=39,column =2,sticky='e')
RGRepresentanteLegalDois = tk.Entry(second_frame, textvariable = RGRepresentanteLegalDois ,width = 20,font=('Arial 10'),borderwidth=5)
RGRepresentanteLegalDois.grid(row=39, column=3,sticky='w')


#-----------------------------------OrgaoEmissorRGDois
OrgaoEmissorRGDois_entry = tk.Label(second_frame, 
         text="Órgão Emissor :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=40,column =2,sticky='e')
OrgaoEmissorRGDois = tk.Entry(second_frame, textvariable = OrgaoEmissorRGDois ,width = 20,font=('Arial 10'),borderwidth=5)
OrgaoEmissorRGDois.grid(row=40, column=3,sticky='w')

#-----------------------------------CargoDois
CargoDois_entry = tk.Label(second_frame, 
         text="Cargo :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=41,column =2,sticky='e')
CargoDois = tk.Entry(second_frame, textvariable = CargoDois ,width = 20,font=('Arial 10'),borderwidth=5)
CargoDois.grid(row=41, column=3,sticky='w')

#-----------------------------------CotaDois
CotaDois = tk.Label(second_frame, 
         text="Cota :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=42,column =2,sticky='e')
CotaDois = tk.Entry(second_frame, textvariable = CotaDois ,width = 20,font=('Arial 10'),borderwidth=5)
CotaDois.grid(row=42, column=3,sticky='w')


#-----------------------------------TelefoneCelularRepresentanteLegalDois
TelefoneCelularRepresentanteLegalDois_entry = tk.Label(second_frame, 
         text="Telefone :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=43,column =2,sticky='e')
TelefoneCelularRepresentanteLegalDois = tk.Entry(second_frame, textvariable = TelefoneCelularRepresentanteLegalDois ,width = 20,font=('Arial 10'),borderwidth=5)
TelefoneCelularRepresentanteLegalDois.grid(row=43, column=3,sticky='w')

#-----------------------------------EmailRepresentanteLegalDois
EmailRepresentanteLegalDois = tk.Label(second_frame, 
         text="Email :",font=('Arial', 10,'bold'),pady=5,padx=10,fg="black", bg="white",
                borderwidth=0).grid(row=44,column =2,sticky='e')
EmailRepresentanteLegalDois = tk.Entry(second_frame, textvariable = EmailRepresentanteLegalDois ,width = 20,font=('Arial 10'),borderwidth=5)
EmailRepresentanteLegalDois.grid(row=44, column=3,sticky='w')


pdf.set(-1)

PDF = Radiobutton(second_frame, text="Gerar PDF's além de arquivos .doc", variable=pdf, value=1,pady=20,fg="black", bg="white")                 
PDF.grid(row=45, column=3,sticky='e')


image = Image.open(caminhoIcone.replace("\Python_ENEL_MD_SE_Base_16092022.docx", "\imgLight.jpeg"))
resize_image = image.resize((70, 30))
img = ImageTk.PhotoImage(resize_image)
label1 = Label(root,image=img , bg='#ff4787')
label1.image = img
label1.pack()

myLabel=Label(root,text="Developed by Yuri Mello",
              font=('Arial', 10,'bold'),padx=20,
              bg='#ff4787',fg="white").pack(side="left",expand=True)

root.config(menu=menubar)
root.mainloop()



